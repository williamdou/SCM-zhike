#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Generate the final 开发与应用报告 as .docx"""

import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

section = doc.sections[0]
section.page_height = Cm(29.7)
section.page_width = Cm(21.0)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(3.18)
section.right_margin = Cm(3.18)

style = doc.styles['Normal']
font = style.font
font.name = 'SimSun'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'SimHei'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
    return h

def add_para(text, bold=False, indent=True):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Pt(24)
    run = p.add_run(text)
    run.font.name = 'SimSun'
    run.font.size = Pt(12)
    run.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.name = 'SimHei'
                run.font.size = Pt(10)
                run.bold = True
                run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx+1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(cell_text)
            run.font.name = 'SimSun'
            run.font.size = Pt(10)
            run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    return table

# ===== TITLE =====
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('开发与应用报告')
run.font.name = 'SimHei'
run.font.size = Pt(22)
run.bold = True
run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('SCM智课——供应链管理研究生智能教学系统')
run.font.name = 'SimHei'
run.font.size = Pt(16)
run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')

doc.add_paragraph()

# Build report content
sections = [
    ('Heading1', '一、开发背景'),
    ('Para', '“供应链管理研究”是面向管理科学与工程、物流工程等专业研究生的核心方法论课程，共48学时，覆盖16个知识模块：从供应链概念、产生逻辑、发展趋势，到战略匹配、整合、网络设计、成员协调，再到产品设计、定价、采购、库存、物流、信息技术、生产计划、风险管理、绩效管理等全过程。'),
    ('Para', '实际教学中面临四个突出痛点：'),
    ('Para', '痛点一：课堂沉闷，学生被动听讲。每章教学陷入固定模式，纯概念讲授印70%课时。课堂主动发言者寥寥（2-3人次/课）。'),
    ('Para', '痛点二：理论与实践脱节。学生知道“牛鞭效应”的定义，却从未体验其在多级供应链中的逐级放大效应；知道“Kraljic矩阵”分类，却从未面对真实供应商数据做出品类策略选择。供应链管理本质上是决策科学，但传统课堂剥夺了学生“做决策”的机会。'),
    ('Para', '痛点三：缺乏动手参与感。国际上SCM教学高度依赖体验式学习——啤酒游戏自1960年代MIT发明以来是教学标配，哈佛商学院大量使用案例决策模拟。但国内研究生课堂受限于技术条件和教师开发能力，“教师讲、学生听”仍是主流。'),
    ('Para', '痛点四：学术视野狭窄。学生对SCM领域学术流派分野、研究脉络演化、热点与方法论要求缺乏系统认知，论文选题缺乏学术根基。'),
    ('Para', '核心矛盾在于：如何在48学时内，既完成16章知识体系教学，又让每章至少有1课时动手实践，还能融合企业案例与学术视野？必须借助生成式人工智能重构教学内容生产方式和课堂组织方式。'),
]

print("Report structure prepared. Now building...")

# Simplify: just write categories to avoid long script issues
# We'll add content in manageable chunks

# Save with basic structure
output_path = r'e:\Claude Education\1教师人工智能应用案例征集材料\开发与应用报告-SCM智课.docx'
doc.save(output_path)
print(f'Saved basic structure to: {output_path}')
